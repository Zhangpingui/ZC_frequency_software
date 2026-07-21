from __future__ import annotations

from io import BytesIO
import json
from pathlib import Path
import re

import pandas as pd

from core.demand_models import ProtectionRule, ProtectionRuleSet


_NUMBER = r"\d+(?:\.\d+)?"
_TOLERANCE = re.compile(
    rf"(?P<center>{_NUMBER})\s*±\s*(?P<tolerance>{_NUMBER})\s*(?P<unit>GHz|MHz|kHz)",
    re.IGNORECASE,
)
_RANGE = re.compile(
    rf"(?P<lower>{_NUMBER})(?:\s*[-–—]\s*(?P<upper>{_NUMBER}))?\s*(?P<unit>GHz|MHz|kHz)",
    re.IGNORECASE,
)
_FREQUENCY_COLUMNS = ("频率", "frequency")
_CATEGORY_COLUMNS = ("用途", "类别", "category", "usage")
_REQUIREMENT_COLUMNS = ("保护要求", "requirement", "保护说明")


def parse_protection_upload(data: bytes, filename: str) -> ProtectionRuleSet:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        entries = _docx_entries(data)
    elif suffix in {".xlsx", ".xls", ".csv"}:
        entries = _tabular_entries(data, suffix)
    elif suffix == ".json":
        entries = _json_entries(data)
    elif suffix == ".txt":
        entries = _text_entries(data)
    else:
        raise ValueError("禁用保护/规则数据仅支持 Word、Excel、CSV、JSON 或 TXT")

    rules: list[ProtectionRule] = []
    warnings: list[str] = []
    for expression, category, requirement in entries:
        parsed = _parse_frequency_expression(expression)
        if not parsed:
            warnings.append(f"无法解析频率表达式：{expression}")
            continue
        for lower, upper in parsed:
            rules.append(ProtectionRule(expression, lower, upper, category, requirement))
    if not rules:
        raise ValueError("未解析到有效的保护规则")
    return ProtectionRuleSet(tuple(rules), filename, tuple(warnings))


def _parse_frequency_expression(expression: str) -> list[tuple[float, float]]:
    value = str(expression).strip()
    if not value:
        return []
    ranges: list[tuple[float, float]] = []
    tolerance_matches = list(_TOLERANCE.finditer(value))
    for match in tolerance_matches:
        factor = _unit_factor(match.group("unit"))
        center = float(match.group("center")) * factor
        tolerance = float(match.group("tolerance")) * factor
        ranges.append((center - tolerance, center + tolerance))
    value_without_tolerances = _TOLERANCE.sub("", value)
    for match in _RANGE.finditer(value_without_tolerances):
        factor = _unit_factor(match.group("unit"))
        lower = float(match.group("lower")) * factor
        upper = float(match.group("upper") or match.group("lower")) * factor
        ranges.append((min(lower, upper), max(lower, upper)))
    return ranges


def _unit_factor(unit: str) -> float:
    return {"ghz": 1000.0, "mhz": 1.0, "khz": 0.001}[unit.lower()]


def _text_entries(data: bytes) -> list[tuple[str, str, str]]:
    entries = []
    for line in data.decode("utf-8-sig", errors="replace").splitlines():
        parts = [part.strip() for part in line.split("|")]
        if not parts or not parts[0] or parts[0] == "频率":
            continue
        entries.append((parts[0], parts[1] if len(parts) > 1 else "", parts[2] if len(parts) > 2 else ""))
    return entries


def _json_entries(data: bytes) -> list[tuple[str, str, str]]:
    payload = json.loads(data.decode("utf-8"))
    if not isinstance(payload, list):
        raise ValueError("JSON 规则数据必须是对象数组")
    entries = []
    for index, item in enumerate(payload, 1):
        if not isinstance(item, dict):
            raise ValueError(f"JSON 第 {index} 项必须是对象")
        expression = item.get("frequency", item.get("频率", ""))
        entries.append((str(expression), str(item.get("category", item.get("用途", ""))), str(item.get("requirement", item.get("保护要求", "")))))
    return entries


def _tabular_entries(data: bytes, suffix: str) -> list[tuple[str, str, str]]:
    try:
        frame = pd.read_csv(BytesIO(data)) if suffix == ".csv" else pd.read_excel(BytesIO(data))
    except ImportError as error:
        raise ValueError("当前环境缺少读取该 Excel 格式所需的组件") from error
    frequency_column = _find_column(frame.columns, _FREQUENCY_COLUMNS)
    if frequency_column is None:
        raise ValueError("规则数据缺少“频率”字段")
    category_column = _find_column(frame.columns, _CATEGORY_COLUMNS)
    requirement_column = _find_column(frame.columns, _REQUIREMENT_COLUMNS)
    return [
        (
            str(row[frequency_column]),
            str(row[category_column]) if category_column and pd.notna(row[category_column]) else "",
            str(row[requirement_column]) if requirement_column and pd.notna(row[requirement_column]) else "",
        )
        for _, row in frame.iterrows()
        if pd.notna(row[frequency_column])
    ]


def _docx_entries(data: bytes) -> list[tuple[str, str, str]]:
    try:
        from docx import Document
    except ImportError as error:
        raise ValueError("当前环境缺少 Word 文档读取组件") from error
    document = Document(BytesIO(data))
    entries = [(paragraph.text, "", "") for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if cells and cells[0] and cells[0] != "频率":
                entries.append((cells[0], cells[1] if len(cells) > 1 else "", cells[2] if len(cells) > 2 else ""))
    return entries


def _find_column(columns: pd.Index, accepted: tuple[str, ...]) -> str | None:
    normalized = {str(column).strip().lower(): str(column) for column in columns}
    for name in accepted:
        if name.lower() in normalized:
            return normalized[name.lower()]
    return None
